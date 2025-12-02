"""
Парсер документов - извлечение текста из различных форматов
Поддерживает: PDF, DOCX, TXT, CSV, XLSX
"""

import io
import csv
from pathlib import Path
from typing import Tuple, Dict, Any, Optional
from datetime import datetime

from preproces.preprocessing import TextPreprocessor


class DocumentParseError(Exception):
    """Ошибка при парсинге документа"""
    pass


class DocumentParser:
    """Парсер для извлечения текста из документов"""
    
    SUPPORTED_FORMATS = {'pdf', 'docx', 'doc', 'txt', 'xlsx', 'xls', 'csv'}
    MAX_FILE_SIZE = 100 * 1024 * 1024  # 100 МБ
    
    @staticmethod
    def get_file_extension(filename: str) -> str:
        """Получить расширение файла"""
        return Path(filename).suffix.lstrip('.').lower()
    
    @staticmethod
    def is_supported_format(filename: str) -> bool:
        """Проверить формат"""
        ext = DocumentParser.get_file_extension(filename)
        return ext in DocumentParser.SUPPORTED_FORMATS
    
    @staticmethod
    async def parse(
        file_content: bytes,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Парсить документ и извлечь текст
        
        Args:
            file_content: Содержимое файла в байтах
            filename: Имя файла
            metadata: Дополнительные метаданные
        
        Returns:
            Кортеж (текст, метаданные)
        
        Raises:
            DocumentParseError: Если не удалось парсить
        """
        
        # Проверка размера
        if len(file_content) > DocumentParser.MAX_FILE_SIZE:
            raise DocumentParseError(
                f"Файл слишком большой ({len(file_content) / 1024 / 1024:.2f} МБ). "
                f"Максимум {DocumentParser.MAX_FILE_SIZE / 1024 / 1024:.0f} МБ"
            )
        
        # Определяем формат
        ext = DocumentParser.get_file_extension(filename)
        
        if not DocumentParser.is_supported_format(filename):
            raise DocumentParseError(
                f"Неподдерживаемый формат: {ext}. "
                f"Поддерживаемые: {', '.join(DocumentParser.SUPPORTED_FORMATS)}"
            )
        
        try:
            if ext == 'pdf':
                text, extracted_metadata = await DocumentParser._parse_pdf(file_content, filename)
            elif ext in ['docx', 'doc']:
                text, extracted_metadata = await DocumentParser._parse_docx(file_content, filename)
            elif ext in ['xlsx', 'xls']:
                text, extracted_metadata = await DocumentParser._parse_excel(file_content, filename)
            elif ext == 'csv':
                text, extracted_metadata = await DocumentParser._parse_csv(file_content, filename)
            elif ext == 'txt':
                text, extracted_metadata = await DocumentParser._parse_text(file_content, filename)
            else:
                raise DocumentParseError(f"Неизвестный формат: {ext}")
            
            # Очищаем текст
            text = TextPreprocessor.clean_text(text)
            
            # Объединяем метаданные
            final_metadata = {
                'filename': filename,
                'format': ext,
                'file_size': len(file_content),
                'parsed_at': datetime.now().isoformat(),
                'char_count': len(text),
                **(metadata or {}),
                **extracted_metadata
            }
            
            return text, final_metadata
            
        except DocumentParseError:
            raise
        except Exception as e:
            raise DocumentParseError(f"Ошибка при парсинге {filename}: {str(e)}")
    
    @staticmethod
    async def _parse_pdf(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Парсить PDF"""
        try:
            import pypdf
            
            pdf_file = io.BytesIO(content)
            reader = pypdf.PdfReader(pdf_file)
            
            text_parts = []
            page_count = len(reader.pages)
            
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"--- Страница {page_num} ---\n{text}")
            
            text = "\n\n".join(text_parts)
            
            metadata = {}
            if reader.metadata:
                metadata = {
                    'title': reader.metadata.get('/Title', ''),
                    'author': reader.metadata.get('/Author', ''),
                    'subject': reader.metadata.get('/Subject', ''),
                }
            
            metadata['page_count'] = page_count
            return text, metadata
            
        except ImportError:
            raise DocumentParseError("pypdf не установлен. Установите: pip install pypdf")
        except Exception as e:
            raise DocumentParseError(f"Ошибка при парсинге PDF: {str(e)}")
    
    @staticmethod
    async def _parse_docx(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Парсить DOCX"""
        try:
            from docx import Document
            
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text_parts = []
            
            # Текст из параграфов
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
            }
            
            try:
                if doc.core_properties:
                    core_props = doc.core_properties
                    if core_props.title:
                        metadata['title'] = core_props.title
                    if core_props.author:
                        metadata['author'] = core_props.author
            except:
                pass
            
            return text, metadata
            
        except ImportError:
            raise DocumentParseError("python-docx не установлен. Установите: pip install python-docx")
        except Exception as e:
            raise DocumentParseError(f"Ошибка при парсинге DOCX: {str(e)}")
    
    @staticmethod
    async def _parse_excel(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Парсить Excel"""
        try:
            from openpyxl import load_workbook
            
            excel_file = io.BytesIO(content)
            workbook = load_workbook(excel_file, data_only=True)
            
            text_parts = []
            sheet_names = workbook.sheetnames
            
            for sheet_name in sheet_names:
                sheet = workbook[sheet_name]
                text_parts.append(f"=== Лист: {sheet_name} ===")
                
                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(v for v in row_values):
                        text_parts.append(" | ".join(row_values))
            
            text = "\n".join(text_parts)
            
            metadata = {
                'sheet_count': len(sheet_names),
                'sheet_names': sheet_names,
            }
            
            return text, metadata
            
        except ImportError:
            raise DocumentParseError("openpyxl не установлен. Установите: pip install openpyxl")
        except Exception as e:
            raise DocumentParseError(f"Ошибка при парсинге Excel: {str(e)}")
    
    @staticmethod
    async def _parse_csv(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Парсить CSV"""
        try:
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            text = None
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    break
                except (UnicodeDecodeError, AttributeError):
                    continue
            
            if text is None:
                raise DocumentParseError("Не удалось определить кодировку")
            
            csv_file = io.StringIO(text)
            reader = csv.reader(csv_file)
            
            text_parts = []
            row_count = 0
            
            for row in reader:
                if row:
                    text_parts.append(" | ".join(str(cell).strip() for cell in row))
                    row_count += 1
            
            text = "\n".join(text_parts)
            
            metadata = {'row_count': row_count}
            return text, metadata
            
        except Exception as e:
            raise DocumentParseError(f"Ошибка при парсинге CSV: {str(e)}")
    
    @staticmethod
    async def _parse_text(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Парсить текстовый файл"""
        try:
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            text = None
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    break
                except (UnicodeDecodeError, AttributeError):
                    continue
            
            if text is None:
                raise DocumentParseError("Не удалось определить кодировку")
            
            metadata = {'line_count': len(text.splitlines())}
            return text, metadata
            
        except Exception as e:
            raise DocumentParseError(f"Ошибка при парсинге текста: {str(e)}")
